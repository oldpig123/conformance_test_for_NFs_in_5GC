import docx
import re
import spacy
import os
from sentence_transformers import SentenceTransformer
from neo4j import GraphDatabase

# configurations
# docx_file_path = "3GPP/33501-j30_new.docx" # Path to the DOCX file
docx_file_path = [
    "3GPP/33501-j30_new.docx",  # Path to the DOCX file
    "3GPP/23502-j40_new.docx"
]

NEO4J_URI = "bolt://localhost:7687"  # Neo4j connection URI
NEO4J_USER = "neo4j"  # Neo4j username
NEO4J_PASSWORD = "12345678"  # Neo4j password

class TGPP_KG_Builder:
    def __init__(self, neo4j_uri=NEO4J_URI, neo4j_user=NEO4J_USER, neo4j_password=NEO4J_PASSWORD):
        print("Initializing 3GPP Knowledge Graph Builder...")
        # Load NLP model
        self.nlp = spacy.load("en_core_web_trf")
        # Load embeddings model
        self.model = SentenceTransformer('Qwen/Qwen3-Embedding-8B')
        # self.model_1 = SentenceTransformer('Linq-AI-Research/Linq-Embed-Mistral')
        # self.model_2 = SentenceTransformer('Qwen/Qwen3-Embedding-8B')
        # Connect to Neo4j database
        self.driver = GraphDatabase.driver(neo4j_uri, auth=(neo4j_user, neo4j_password))
        print("Initialization complete.")

    def close(self):
        """Close the Neo4j driver connection."""
        if self.driver:
            self.driver.close()
            print("Neo4j connection closed.")
            
    # def extract_requirements(self,file_path):
    #     """Extract requirements from the DOCX file."""
    #     try:
    #         doc = docx.Document(file_path)
    #         full_text = '\n'.join([para.text for para in doc.paragraphs])
    #         nlp_doc = self.nlp(full_text)
            
    #         # keyword patterns for requirements
    #         req_patterns = r'\b(?:shall|must|should|may|can)\b'
    #         requirements = [sent.text.strip() for sent in nlp_doc.sents if re.search(req_patterns, sent.text, re.IGNORECASE)]
            
    #         print(f"Extracted {len(requirements)} requirements.")
    #         return requirements
    #     except Exception as e:
    #         print(f"Error reading DOCX file: {e}")
    #         return []
    
    def extract_requirements(self, file_path):
        """extract requirements from the DOCX file."""
        print(f"Extracting sections from {file_path}...")
        try:
            doc = docx.Document(file_path)
            sections = []
            current_heading = None
            current_text = ""
            
            for para in doc.paragraphs:
                # check if the  paragraph is a heading
                if para.style.name.startswith('Heading'):
                    # if we have collected text for a previous heading, save it
                    if current_heading and current_text.strip():
                        sections.append({'title' : current_heading, 'text': current_text.strip()})
                        
                    # start a new section
                    current_heading = para.text
                    current_text = ""
                else:
                    # append praragraph text to the current section
                    current_text += para.text + "\n"
            
            # append the lase section after the loop
            if current_text.strip():
                sections.append({'title': current_heading, 'text': current_text.strip()})
            
            print( f"Extracted {len(sections)} sections from the document.")
            return sections
        
        except Exception as e:
            print(f"Error reading DOCX file: {e}")
            return []

    # def process_requirements(self, requirements):
    #     """Gererate embeddings and process requirements."""
    #     # 1. Generate embeddings
    #     embedding = self.model.encode(requirements, convert_to_tensor=True).tolist()
        
    #     # 2. Etract entities and relations
    #     doc = self.nlp(requirements)
    #     subject = None
    #     obj = None
    #     # find the nominal subject and direct object of the root verb
    #     for token in doc:
    #         if "subj" in token.dep_ and not subject:
    #             subject = token.text
    #         if "obj" in token.dep_ and not obj:
    #             obj = token.text
                
    #     # Fallback if parsing is tricky
    #     if not subject: 
    #         subject = "Unidentified"
    #     if not obj:
    #         obj = "Unidentified"
            
    #     return {
    #         "text": requirements,
    #         "embedding": embedding,
    #         "subject": subject,
    #         "object": obj
    #     }
        
    def process_requirements(self, requirement,source_file):
        """Generate embeddings and extract entitis for a whole section."""
        # # 1.Gernerate a single embedding for the entire section text
        # embedding = self.model.encode(requirement['text'], convert_to_tensor=True).tolist()
        
        # # 2. Extract unique extities (nouns ans proper nouns) from the text
        # doc = self.nlp(requirement['text'])
        # entities = list(set([chunck.text for chunck in doc.noun_chunks if len(chunck.text.split()) < 4]))
        
        # return {
        #     "title": requirement['title'],
        #     "text": requirement['text'],
        #     "embedding": embedding,
        #     "entities": entities
        # }

        embedding = self.model.encode(requirement['text'], convert_to_tensor=True).tolist()
        # embedding_1 = self.model_1.encode(requirement['text'], convert_to_tensor=True).tolist()
        # embedding_2 = self.model_2.encode(requirement['text'], convert_to_tensor=True).tolist()
        doc = self.nlp(requirement['text'])
        entities = list(set([chunk.text for chunk in doc.noun_chunks if len(chunk.text.split()) < 4]))
        
        return {
            "title": requirement['title'],
            "text": requirement['text'],
            "embedding": embedding,
            # "embedding_1": embedding_1,
            # "embedding_2": embedding_2,
            "entities": entities,
            "source": os.path.basename(source_file)  # Add source file name
        }
    
    @staticmethod
    def _create_requirment_nodes(tx,procesed_data):
        # """The transaction funciton to create nodes and relationships."""
        # query = """
        # UNWIND $data as item
        # // Create the requirement node
        # CREATE (r:Requirement {text: item.text, embedding: item.embedding})
        
        # // merge the subject and object entities
        # MERGE (subj:Entity {name: item.subject})
        # MERGE (obj:Entity {name: item.object})
        
        # // create relationships
        # MERGE (r)-[:DEFINES_SUBJECT]->(subj)
        # MERGE (r)-[:DEFINES_OBJECT]->(obj)
        # """
        
        """ the transaction function to create section nodes and their entity relationships."""
        # query = """
        # UNWIND $data as item
        # // Create the Section node
        # CREATE (s:Requirement {title: item.title, text: item.text, embedding: item.embedding})
        
        # // For each entity found in this section, create an Entity node and a relationship
        # FOREACH (entity_name IN item.entities |
        #     MERGE (e:Entity {name: entity_name})
        #     MERGE (s)-[:CONTAINS_ENTITY]->(e)
        # )
        # """
        query = """
        UNWIND $data as item
        // Merge the Document node to avoid duplicates
        MERGE (d:Document {name: item.source})
        
        // Create the Section node
        CREATE (s:Requirement {title: item.title, text: item.text, embedding: item.embedding})
        
        // Link the Section to its source Document
        MERGE (s)-[:PART_OF]->(d)
        
        // For each entity, create it and link it to the section
        FOREACH (entity_name IN item.entities |
            MERGE (e:Entity {name: entity_name})
            MERGE (s)-[:CONTAINS_ENTITY]->(e)
        )
        """
        
        tx.run(query, data=procesed_data)
        

    def populate_graph(self, processed_data):
        """write the prcessed data to the neo4j database."""
        with self.driver.session(database="neo4j") as session:
            session.execute_write(self._create_requirment_nodes, processed_data)
            
    def run_pipline(self, list_of_docx_file_path):
        """Execute the full pipeline from parsing to graph population."""
        # # step 1: Extract requirements
        # requirements = self.extract_requirements(docx_file_path)
        # if not requirements:
        #     print("No requirments found. Exiting.")
        #     return
        
        # # step 2: process each requirment to get embeddings and entities
        # processed_data = [self.process_requirements(req) for req in requirements]
        
        all_processed_data = []
        
        for docx_file in list_of_docx_file_path:
            # step 1: extract requirements
            requirements = self.extract_requirements(docx_file)
            if not requirements:
                print(f"No requirements found in {docx_file}. Skipping.")
                continue
            
            # step 2: process each requirement and passing the source file path
            proceccessed_requirements = [self.process_requirements(req, docx_file) for req in requirements]
            all_processed_data.extend(proceccessed_requirements)
        
        # step 3: populate the knowledge graph
        # print(f"Populating graph with {len(processed_data)} requirements...")
        # self.populate_graph(processed_data)
        # print("pipeline complete successfully. knowledge graph been updated.")
        
        # step 3: populate the knowledge graph with requirements from all files
        if all_processed_data:
            print(f"Populating graph with {len(all_processed_data)} requirements from {len(list_of_docx_file_path)} files...")
            self.populate_graph(all_processed_data)
            print("Pipeline completed successfully. Knowledge graph has been updated.")
        
if __name__ == "__main__":
    # create and instance of the builder
    builder = TGPP_KG_Builder(NEO4J_URI, NEO4J_USER, NEO4J_PASSWORD)

    # run the entire process
    builder.run_pipline(docx_file_path)
    
    # close the connection
    builder.close()